"""Tests for TODO markdown / docx export."""

import os
import tempfile
import unittest

from app.model.annotations import Annotation, KIND_COMMENT, KIND_TEXTBOX
from app.export.todo_export import (
    export_markdown, export_docx, GROUP_PAGE, GROUP_SHEET,
)


class TestTodoExport(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.mkdtemp()
        self.todos = [
            Annotation(page=4, kind=KIND_COMMENT, text="Verify breaker",
                       author="Eli", is_todo=True, created="2026-06-21T10:00:00"),
            Annotation(page=4, kind=KIND_TEXTBOX, text="Revision B",
                       author="Eli", is_todo=True, todo_done=True,
                       created="2026-06-21T11:00:00"),
            Annotation(page=6, kind=KIND_COMMENT, text="Wire mislabeled",
                       author="Bob", is_todo=True, created="2026-06-21T12:00:00"),
        ]
        self.sheets = {4: "300", 6: "601"}

    def test_markdown_by_page(self):
        p = export_markdown(self.todos, os.path.join(self.tmp, "t.md"),
                            group_by=GROUP_PAGE)
        text = open(p).read()
        self.assertIn("## Page 5", text)
        self.assertIn("## Page 7", text)
        self.assertIn("- [ ] Verify breaker — p.5, Eli", text)
        self.assertIn("- [x] Revision B", text)  # done item

    def test_markdown_by_sheet(self):
        p = export_markdown(self.todos, os.path.join(self.tmp, "s.md"),
                            group_by=GROUP_SHEET, sheet_labels=self.sheets)
        text = open(p).read()
        self.assertIn("## Sheet 300", text)
        self.assertIn("## Sheet 601", text)
        # the per-row meta carries the sheet too
        self.assertIn("p.5, sheet 300", text)

    def test_docx_builds(self):
        from docx import Document as Docx
        p = export_docx(self.todos, os.path.join(self.tmp, "t.docx"),
                        group_by=GROUP_PAGE, sheet_labels=self.sheets)
        self.assertTrue(os.path.getsize(p) > 0)
        d = Docx(p)
        # at least the two page tables exist
        self.assertGreaterEqual(len(d.tables), 2)
        # header row present, now incl. a Sheet column
        self.assertEqual(d.tables[0].rows[0].cells[0].text, "Done")
        self.assertEqual(d.tables[0].rows[0].cells[3].text, "Sheet")


if __name__ == "__main__":
    unittest.main()
