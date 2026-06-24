"""Turn captured PDF regions into useful documents.

Each region is classified (by Claude vision, see
:func:`app.extraction.claude_api.extract_region_content`) as either a **table**
or a **body of prose**.  This module materialises a list of such region results
into the chosen outputs:

* **Excel** (``.xlsx``) - one worksheet per table region (rebuilt grid).
* **Word** (``.docx``) - prose regions as headed sections (good for editing).
* **Markdown** (``.md``) - *everything* in one file (tables as markdown tables,
  prose as text) - the no-routing "just give me the text" option.

A region result is a plain dict::

    {"page": int, "index": int, "type": "table"|"text",
     "title": str, "rows": [[cell, ...], ...], "text": str}

GUI-free and unit-testable; the writers degrade to no-ops when a region set has
nothing of that kind.
"""

from __future__ import annotations

import os
import re
from typing import Iterable, Optional

TYPE_TABLE = "table"
TYPE_TEXT = "text"


def region_title(r: dict) -> str:
    """Human caption for a region (its AI title, else page/index)."""
    t = (r.get("title") or "").strip()
    if t:
        return t
    return f"Page {r.get('page', 0) + 1} · region {r.get('index', 0) + 1}"


def has_tables(results: Iterable[dict]) -> bool:
    return any(r.get("type") == TYPE_TABLE and r.get("rows") for r in results)


def has_text(results: Iterable[dict]) -> bool:
    return any(r.get("type") != TYPE_TABLE and (r.get("text") or "").strip()
               for r in results)


# --- Excel (tables) ---------------------------------------------------------


def _safe_sheet_name(name: str, used: set) -> str:
    name = re.sub(r"[\[\]:*?/\\]", " ", name).strip() or "Table"
    name = name[:31]
    base = name
    n = 2
    while name.lower() in used:
        suffix = f" ({n})"
        name = base[:31 - len(suffix)] + suffix
        n += 1
    used.add(name.lower())
    return name


def write_excel(results: Iterable[dict], path: str) -> Optional[str]:
    """Write each table region to its own worksheet.  Returns ``path`` or
    ``None`` when there are no tables."""
    results = list(results)
    if not has_tables(results):
        return None
    from openpyxl import Workbook
    wb = Workbook()
    wb.remove(wb.active)
    used: set = set()
    for r in results:
        if r.get("type") != TYPE_TABLE or not r.get("rows"):
            continue
        ws = wb.create_sheet(_safe_sheet_name(region_title(r), used))
        for row in r["rows"]:
            ws.append(list(row) if isinstance(row, (list, tuple)) else [row])
    if not wb.sheetnames:                       # safety: nothing got written
        return None
    wb.save(path)
    return path


# --- Word (prose) -----------------------------------------------------------


def write_word(results: Iterable[dict], path: str) -> Optional[str]:
    """Write each prose region as a headed section.  Returns ``path`` or
    ``None`` when there is no prose."""
    results = list(results)
    if not has_text(results):
        return None
    from docx import Document as Docx
    doc = Docx()
    first = True
    for r in results:
        if r.get("type") == TYPE_TABLE or not (r.get("text") or "").strip():
            continue
        if not first:
            doc.add_paragraph("")
        first = False
        doc.add_heading(region_title(r), level=2)
        for line in r["text"].split("\n"):
            doc.add_paragraph(line)
    doc.save(path)
    return path


# --- Markdown (everything) --------------------------------------------------


def _md_escape(cell: str) -> str:
    return str(cell).replace("|", r"\|").replace("\n", " ").strip()


def _table_to_md(rows) -> str:
    rows = [list(r) if isinstance(r, (list, tuple)) else [r] for r in rows if r]
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    header, *body = rows
    out = ["| " + " | ".join(_md_escape(c) for c in header) + " |",
           "| " + " | ".join("---" for _ in header) + " |"]
    for r in body:
        out.append("| " + " | ".join(_md_escape(c) for c in r) + " |")
    return "\n".join(out)


def write_markdown(results: Iterable[dict], path: str) -> Optional[str]:
    """Write *all* regions into one markdown file (tables as markdown tables,
    prose as text)."""
    results = list(results)
    if not results:
        return None
    blocks = []
    for r in results:
        blocks.append(f"## {region_title(r)}")
        if r.get("type") == TYPE_TABLE and r.get("rows"):
            blocks.append(_table_to_md(r["rows"]))
        else:
            blocks.append((r.get("text") or "").strip() or "_(no text captured)_")
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("\n\n".join(blocks) + "\n")
    return path


# --- orchestrator -----------------------------------------------------------


def export_regions(results: Iterable[dict], out_dir: str, base: str,
                   formats: Iterable[str]) -> list:
    """Write the requested ``formats`` (``"excel"``/``"word"``/``"markdown"``).

    Returns the list of paths actually written (a format that has no matching
    content is skipped).
    """
    results = list(results)
    formats = set(formats)
    os.makedirs(out_dir, exist_ok=True)
    written: list = []
    if "excel" in formats:
        p = write_excel(results, os.path.join(out_dir, f"{base}_tables.xlsx"))
        if p:
            written.append(p)
    if "word" in formats:
        p = write_word(results, os.path.join(out_dir, f"{base}_text.docx"))
        if p:
            written.append(p)
    if "markdown" in formats:
        p = write_markdown(results, os.path.join(out_dir, f"{base}_capture.md"))
        if p:
            written.append(p)
    return written
