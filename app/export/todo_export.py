"""TODO list export to Markdown and DOCX (Phase 5).

GUI-free; operates on :class:`app.model.annotations.Annotation` objects flagged
as TODO items.
"""

from __future__ import annotations

from datetime import datetime
from typing import Iterable, Optional

GROUP_PAGE = "page"          # group by PDF page (header "Page N")
GROUP_SHEET = "sheet"        # group by drawing sheet number (header "Sheet <n>")
GROUP_COMMENTER = "commenter"
GROUP_NONE = "none"


def _date_only(iso: str) -> str:
    if not iso:
        return ""
    try:
        return datetime.fromisoformat(iso).strftime("%Y-%m-%d")
    except ValueError:
        return iso[:10]


def _sheet_of(ann, sheet_labels: Optional[dict]) -> str:
    return (sheet_labels or {}).get(ann.page, "") if sheet_labels else ""


def _group_key(ann, group_by: str, sheet_labels: Optional[dict] = None):
    if group_by == GROUP_PAGE:
        return f"Page {ann.page + 1}"
    if group_by == GROUP_SHEET:
        lbl = _sheet_of(ann, sheet_labels)
        return f"Sheet {lbl}" if lbl else "(no sheet)"
    if group_by == GROUP_COMMENTER:
        return ann.author or "(unknown)"
    return ""


def _grouped(todos: list, group_by: str, sheet_labels: Optional[dict] = None) -> list:
    """Return ``[(heading, [anns...]), ...]`` preserving first-seen order."""
    if group_by == GROUP_NONE:
        return [("", list(todos))]
    order: list = []
    buckets: dict = {}
    for a in todos:
        k = _group_key(a, group_by, sheet_labels)
        if k not in buckets:
            buckets[k] = []
            order.append(k)
        buckets[k].append(a)
    return [(k, buckets[k]) for k in order]


def export_markdown(todos: Iterable, path: str, group_by: str = GROUP_PAGE,
                    title: str = "TODO List", sheet_labels: Optional[dict] = None) -> str:
    """GitHub task-list markdown grouped by page/sheet/commenter with ``##`` headers."""
    todos = list(todos)
    lines = [f"# {title}", "",
             f"_Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}_", ""]
    for heading, group in _grouped(todos, group_by, sheet_labels):
        if heading:
            lines.append(f"## {heading}")
            lines.append("")
        for a in group:
            box = "x" if a.todo_done else " "
            text = " ".join((a.text or a.snippet()).split())
            meta = f"p.{a.page + 1}"
            sheet = _sheet_of(a, sheet_labels)
            if sheet:
                meta += f", sheet {sheet}"
            if a.author:
                meta += f", {a.author}"
            d = _date_only(a.created)
            if d:
                meta += f", {d}"
            tag = f" #{','.join(a.tags)}" if a.tags else ""
            lines.append(f"- [{box}] {text} — {meta}{tag}")
        lines.append("")
    content = "\n".join(lines).rstrip() + "\n"
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(content)
    return path


def export_docx(todos: Iterable, path: str, group_by: str = GROUP_PAGE,
                title: str = "TODO List", sheet_labels: Optional[dict] = None) -> str:
    """DOCX table: Done | Text | Page | Sheet | Commenter | Date with ☐/☑ glyphs."""
    from docx import Document as Docx
    from docx.shared import Pt

    todos = list(todos)
    doc = Docx()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    for heading, group in _grouped(todos, group_by, sheet_labels):
        if heading:
            doc.add_heading(heading, level=2)
        table = doc.add_table(rows=1, cols=6)
        try:
            table.style = "Light Grid Accent 1"
        except Exception:
            pass
        hdr = table.rows[0].cells
        for i, name in enumerate(["Done", "Text", "Page", "Sheet", "Commenter", "Date"]):
            hdr[i].text = name
            for p in hdr[i].paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.size = Pt(10)
        for a in group:
            cells = table.add_row().cells
            cells[0].text = "☑" if a.todo_done else "☐"  # ☑ / ☐
            cells[1].text = " ".join((a.text or a.snippet()).split())
            cells[2].text = str(a.page + 1)
            cells[3].text = _sheet_of(a, sheet_labels)
            cells[4].text = a.author or ""
            cells[5].text = _date_only(a.created)
    doc.save(path)
    return path
